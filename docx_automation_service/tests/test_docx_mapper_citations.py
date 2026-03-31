from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from docx_automation_service.core.models import Chunk, ChunkRef
from docx_automation_service.services.docx_mapper import DocxMapper


def test_inline_citation_rendered_as_superscript_outside_references() -> None:
    doc = Document()
    doc.add_paragraph("原文")

    mapper = DocxMapper()
    chunk = Chunk(
        chunk_id="p-0",
        text="原文",
        ref=ChunkRef(block_type="paragraph", paragraph_index=0, is_reference_section=False),
    )
    mapper.apply_text(doc, chunk, "测试内容[12]继续")

    para = doc.paragraphs[0]
    citation_runs = [r for r in para.runs if r.text == "[12]"]
    assert citation_runs
    assert citation_runs[0].font.superscript is True


def test_inline_citation_not_superscript_in_references_section() -> None:
    doc = Document()
    doc.add_paragraph("参考文献")

    mapper = DocxMapper()
    chunk = Chunk(
        chunk_id="p-0",
        text="参考文献",
        ref=ChunkRef(block_type="paragraph", paragraph_index=0, is_reference_section=True),
    )
    mapper.apply_text(doc, chunk, "[12] 作者. 题目")

    para = doc.paragraphs[0]
    citation_runs = [r for r in para.runs if "[12]" in r.text]
    assert citation_runs
    assert citation_runs[0].font.superscript is not True


def test_extract_chunks_marks_reference_section() -> None:
    doc = Document()
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("这里是正文[1]")
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 张三. 论文A")

    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    with TemporaryDirectory() as tmp_dir:
        docx_path = Path(tmp_dir) / "sample.docx"
        docx_path.write_bytes(data)
        mapper = DocxMapper()
        _, chunks = mapper.extract_chunks(docx_path)

    assert chunks[1].ref.is_reference_section is False
    assert chunks[2].ref.is_reference_section is True
    assert chunks[3].ref.is_reference_section is True
