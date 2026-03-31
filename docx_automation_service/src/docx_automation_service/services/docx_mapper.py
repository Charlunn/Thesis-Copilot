from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from docx_automation_service.core.models import Chunk, ChunkRef
from docx_automation_service.services.text_guard import is_heading_like

_INLINE_CITATION_RE = re.compile(r"(\[[1-9]\d{0,2}\])")
_REFERENCE_HEADING_RE = re.compile(r"^(参考文献|references?)$", re.IGNORECASE)
_NON_REFERENCE_SECTION_HEADING_RE = re.compile(
    r"^(附录|致谢|acknowledg(?:e)?ments?|appendix|作者简介|声明|结论|总结)$",
    re.IGNORECASE,
)


class DocxMapper:
    def extract_chunks(self, file_path: Path) -> tuple[Document, list[Chunk]]:
        doc = Document(str(file_path))
        chunks: list[Chunk] = []

        p_idx = 0
        in_references = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style is not None else None
                if _REFERENCE_HEADING_RE.match(text):
                    in_references = True
                elif in_references and is_heading_like(text, style_name) and _NON_REFERENCE_SECTION_HEADING_RE.match(text):
                    in_references = False

                chunks.append(
                    Chunk(
                        chunk_id=f"p-{p_idx}",
                        text=text,
                        ref=ChunkRef(
                            block_type="paragraph",
                            paragraph_index=p_idx,
                            style_name=style_name,
                            is_heading=is_heading_like(text, style_name),
                            is_reference_section=in_references,
                        ),
                    )
                )
            p_idx += 1

        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = "\n".join(p.text for p in cell.paragraphs).strip()
                    if not text:
                        continue
                    chunks.append(
                        Chunk(
                            chunk_id=f"t-{t_idx}-r-{r_idx}-c-{c_idx}",
                            text=text,
                            ref=ChunkRef(
                                block_type="table_cell",
                                table_index=t_idx,
                                row_index=r_idx,
                                cell_index=c_idx,
                                is_heading=False,
                                is_reference_section=False,
                            ),
                        )
                    )

        return doc, chunks

    def apply_text(self, doc: Document, chunk: Chunk, new_text: str) -> None:
        ref = chunk.ref
        text = new_text.strip()
        if not text:
            return

        if ref.block_type == "paragraph" and ref.paragraph_index is not None:
            para = doc.paragraphs[ref.paragraph_index]
            self._replace_in_paragraph(para, text, allow_inline_citation_superscript=not ref.is_reference_section)
            return

        if (
            ref.block_type == "table_cell"
            and ref.table_index is not None
            and ref.row_index is not None
            and ref.cell_index is not None
        ):
            cell = doc.tables[ref.table_index].rows[ref.row_index].cells[ref.cell_index]
            first_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            self._replace_in_paragraph(first_para, text, allow_inline_citation_superscript=True)
            for extra_para in cell.paragraphs[1:]:
                self._replace_in_paragraph(extra_para, "", allow_inline_citation_superscript=False)

    @staticmethod
    def _replace_in_paragraph(para, text: str, *, allow_inline_citation_superscript: bool) -> None:
        for run in para.runs:
            run.text = ""

        if not allow_inline_citation_superscript:
            para.add_run(text)
            return

        parts = _INLINE_CITATION_RE.split(text)
        for part in parts:
            if not part:
                continue
            run = para.add_run(part)
            if _INLINE_CITATION_RE.fullmatch(part):
                run.font.superscript = True
