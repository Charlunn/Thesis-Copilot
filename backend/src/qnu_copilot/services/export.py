from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Protocol

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pydantic import BaseModel, Field

from qnu_copilot.domain.enums import GenericStatus, WorkflowStage
from qnu_copilot.domain.models import BibtexEntry, ExportHistoryItem, ProjectState
from qnu_copilot.services.errors import ConflictError
from qnu_copilot.services.filesystem import sanitize_title
from qnu_copilot.services.workspace import WorkspaceManager


CITATION_RE = re.compile(r"【文献\s*0*(\d+)】")


class ExportReferenceItem(BaseModel):
    effective_index: int
    title: str
    bibtex_key: str | None = None


class DocumentExportInput(BaseModel):
    document_title: str
    final_blocks: list[dict[str, Any]] = Field(default_factory=list)
    references: list[ExportReferenceItem] = Field(default_factory=list)
    citation_mapping: dict[str, int] = Field(default_factory=dict)
    bibtex_entries: list[BibtexEntry] = Field(default_factory=list)
    template_path: str = ""
    output_path: str


class DocumentExportResult(BaseModel):
    status: str
    message: str
    output_path: str


class ProjectExportResult(BaseModel):
    output_path: str
    log_path: str
    reference_count: int
    message: str


class DocumentExportService(Protocol):
    def export(self, export_input: DocumentExportInput) -> DocumentExportResult:
        ...


class LocalDocumentExportService:
    def export(self, export_input: DocumentExportInput) -> DocumentExportResult:
        template_path = Path(export_input.template_path) if export_input.template_path else None
        if template_path and template_path.exists():
            document = Document(str(template_path))
        else:
            document = Document()

        self._write_title(document, export_input.document_title)
        for block in export_input.final_blocks:
            self._write_block(document, block, export_input.citation_mapping)

        references = self._build_reference_lines(
            export_input.references,
            export_input.bibtex_entries,
        )
        if references:
            document.add_page_break()
            document.add_heading("参考文献", level=1)
            for reference_line in references:
                paragraph = document.add_paragraph()
                paragraph.style = document.styles["Normal"]
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.add_run(reference_line)

        output_path = Path(export_input.output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return DocumentExportResult(
            status="completed",
            message="Document exported successfully.",
            output_path=str(output_path.resolve()),
        )

    def validate_block_content(
        self,
        final_blocks: list[dict[str, Any]],
        citation_mapping: dict[str, int],
    ) -> list[str]:
        """Validate block content for issues like missing citations or empty blocks."""
        issues: list[str] = []
        max_citation = max(citation_mapping.values()) if citation_mapping else 0

        for idx, block in enumerate(final_blocks, start=1):
            block_title = block.get("title", f"第{idx}块")
            content = block.get("content", [])
            
            if not content:
                issues.append(f"【{block_title}】内容为空")
                continue

            has_text = False
            has_heading = False
            
            for element in content:
                element_type = element.get("type")
                text = element.get("text", "")
                
                if element_type in ("h1", "h2", "h3"):
                    has_heading = True
                    if not text.strip():
                        issues.append(f"【{block_title}】标题为空")
                
                if element_type in ("p", "list"):
                    if text.strip():
                        has_text = True
                        # Check for citation references
                        citations = CITATION_RE.findall(text)
                        for citation in citations:
                            citation_num = int(citation.lstrip("0") or "0")
                            if citation_num > max_citation:
                                issues.append(f"【{block_title}】引用编号 [{citation_num}] 超出有效文献数量 ({max_citation})")
            
            if not has_text:
                issues.append(f"【{block_title}】无实际文本内容")
            
            if not has_heading:
                issues.append(f"【{block_title}】缺少标题")

        return issues

    def validate_structure(
        self,
        final_blocks: list[dict[str, Any]],
    ) -> list[str]:
        """Validate structural integrity of the document."""
        issues: list[str] = []
        
        if not final_blocks:
            issues.append("文档没有任何正文块")
            return issues

        has_first_level_heading = False
        heading_order: list[tuple[int, str, int]] = []  # (level, title, block_index)
        
        for idx, block in enumerate(final_blocks, start=1):
            content = block.get("content", [])
            block_has_h1 = False
            
            for element in content:
                element_type = element.get("type")
                text = element.get("text", "").strip()
                
                if element_type == "h1":
                    has_first_level_heading = True
                    block_has_h1 = True
                    heading_order.append((1, text, idx))
                elif element_type == "h2":
                    heading_order.append((2, text, idx))
                elif element_type == "h3":
                    heading_order.append((3, text, idx))

            if block_has_h1 and idx > 1:
                # First h1 should be early in the document
                if not has_first_level_heading:
                    issues.append(f"第{idx}块是第一级标题，但文档前面缺少一级标题")

        if not has_first_level_heading:
            issues.append("文档缺少一级标题（如「引言」「总结与展望」等）")

        # Check heading sequence
        expected_level = 1
        for level, title, block_idx in heading_order:
            if level > expected_level + 1:
                issues.append(f"第{block_idx}块标题层级跳跃：从{expected_level}级直接到{level}级")
            expected_level = level

        return issues

    def _write_title(self, document: Document, title: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title.strip())
        run.bold = True
        run.font.size = Pt(18)

    def _write_block(
        self,
        document: Document,
        block: dict[str, Any],
        citation_mapping: dict[str, int],
    ) -> None:
        for element in block.get("content", []):
            element_type = element.get("type")
            if element_type == "h1":
                paragraph = document.add_heading(level=1)
                self._append_rich_text(paragraph, element.get("text", ""), citation_mapping)
                continue
            if element_type == "h2":
                paragraph = document.add_heading(level=2)
                self._append_rich_text(paragraph, element.get("text", ""), citation_mapping)
                continue
            if element_type == "h3":
                paragraph = document.add_heading(level=3)
                self._append_rich_text(paragraph, element.get("text", ""), citation_mapping)
                continue
            if element_type == "p":
                paragraph = document.add_paragraph()
                self._append_rich_text(paragraph, element.get("text", ""), citation_mapping)
                continue
            if element_type == "list":
                for item in element.get("items", []):
                    paragraph = document.add_paragraph(style="List Bullet")
                    self._append_rich_text(paragraph, item, citation_mapping)
                continue
            if element_type == "table_placeholder":
                paragraph = document.add_paragraph()
                paragraph.add_run("[表格占位] ")
                self._append_rich_text(paragraph, element.get("text", ""), citation_mapping)

    def _append_rich_text(
        self,
        paragraph,
        text: str,
        citation_mapping: dict[str, int],
    ) -> None:
        cursor = 0
        for match in CITATION_RE.finditer(text):
            if match.start() > cursor:
                paragraph.add_run(text[cursor : match.start()])
            source_number = match.group(1).lstrip("0") or "0"
            mapped_number = citation_mapping.get(source_number) or citation_mapping.get(
                source_number.zfill(2)
            )
            citation_run = paragraph.add_run(f"[{mapped_number or source_number}]")
            citation_run.font.superscript = True
            cursor = match.end()
        if cursor < len(text):
            paragraph.add_run(text[cursor:])

    def _build_reference_lines(
        self,
        references: list[ExportReferenceItem],
        bibtex_entries: list[BibtexEntry],
    ) -> list[str]:
        entry_by_key = {entry.key: entry for entry in bibtex_entries if entry.key}
        entry_by_title = {
            sanitize_title(entry.title or ""): entry
            for entry in bibtex_entries
            if entry.title
        }
        lines: list[str] = []
        for reference in references:
            matched = entry_by_key.get(reference.bibtex_key) or entry_by_title.get(
                sanitize_title(reference.title)
            )
            formatted = self._format_reference_line(reference.effective_index, reference.title, matched)
            lines.append(formatted)
        return lines

    def _format_reference_line(
        self,
        effective_index: int,
        fallback_title: str,
        bibtex_entry: BibtexEntry | None,
    ) -> str:
        if not bibtex_entry:
            return f"[{effective_index}] {fallback_title}."

        raw_text = bibtex_entry.raw_text
        title = self._extract_bibtex_field(raw_text, "title") or fallback_title
        author = self._extract_bibtex_field(raw_text, "author")
        journal = self._extract_bibtex_field(raw_text, "journal")
        booktitle = self._extract_bibtex_field(raw_text, "booktitle")
        publisher = self._extract_bibtex_field(raw_text, "publisher")
        year = self._extract_bibtex_field(raw_text, "year")

        source = journal or booktitle or publisher
        parts = [f"[{effective_index}]"]
        if author:
            parts.append(self._normalize_author(author))
        parts.append(f"{title}.")
        if source:
            parts.append(source)
        if year:
            parts.append(year)
        return " ".join(part.strip() for part in parts if part and part.strip())

    def _normalize_author(self, author_text: str) -> str:
        collapsed = re.sub(r"\s+", " ", author_text.replace("\n", " ")).strip()
        return collapsed.replace(" and ", ", ")

    def _extract_bibtex_field(self, bibtex_text: str, field_name: str) -> str | None:
        field_match = re.search(
            rf"{field_name}\s*=\s*([{{\"])",
            bibtex_text,
            flags=re.IGNORECASE,
        )
        if not field_match:
            return None

        opening = field_match.group(1)
        cursor = field_match.end()

        if opening == '"':
            collected: list[str] = []
            escaped = False
            while cursor < len(bibtex_text):
                current = bibtex_text[cursor]
                cursor += 1
                if escaped:
                    collected.append(current)
                    escaped = False
                    continue
                if current == "\\":
                    escaped = True
                    continue
                if current == '"':
                    break
                collected.append(current)
            value = "".join(collected).strip()
            return value or None

        depth = 1
        collected = []
        while cursor < len(bibtex_text) and depth > 0:
            current = bibtex_text[cursor]
            cursor += 1
            if current == "{":
                depth += 1
                collected.append(current)
                continue
            if current == "}":
                depth -= 1
                if depth == 0:
                    break
                collected.append(current)
                continue
            collected.append(current)

        value = "".join(collected).strip()
        return value or None


class ProjectExportService:
    def __init__(
        self,
        workspace_manager: WorkspaceManager,
        document_export_service: DocumentExportService,
    ) -> None:
        self.workspace_manager = workspace_manager
        self.document_export_service = document_export_service

    def export_project(
        self,
        project_id: str,
        *,
        output_filename: str | None = None,
    ) -> ProjectExportResult:
        state = self.workspace_manager.load_state(project_id)
        self._validate_export_ready(state)

        output_dir = self.workspace_manager.get_project_root(project_id) / "output"
        output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        safe_name = sanitize_title(output_filename or state.project.title) or "qnu_thesis"
        output_path = output_dir / f"{timestamp}_{safe_name}.docx"
        log_path = output_dir / f"{timestamp}_{safe_name}.export.json"
        template_path = self._resolve_template_path(state.template_id)

        final_blocks = [
            block.normalized_json
            for block in sorted(state.generation.blocks, key=lambda item: item.block_index)
            if block.normalized_json is not None
        ]
        abstract_block = self._build_abstract_block(state)
        if abstract_block is not None:
            final_blocks = [abstract_block, *final_blocks]
        citation_mapping = self._build_citation_mapping(state)

        # Perform validation before export
        if isinstance(self.document_export_service, LocalDocumentExportService):
            content_issues = self.document_export_service.validate_block_content(
                final_blocks, citation_mapping
            )
            structure_issues = self.document_export_service.validate_structure(final_blocks)
            
            all_issues = content_issues + structure_issues
            if all_issues:
                issues_text = "; ".join(all_issues)
                raise ConflictError(f"导出前检查发现问题：{issues_text}")

        export_input = DocumentExportInput(
            document_title=state.project.title,
            final_blocks=final_blocks,
            references=[
                ExportReferenceItem(
                    effective_index=item.effective_index,
                    title=item.title,
                    bibtex_key=item.bibtex_key,
                )
                for item in sorted(
                    state.references.processed_items,
                    key=lambda item: item.effective_index,
                )
            ],
            citation_mapping=citation_mapping,
            bibtex_entries=state.references.bibtex_entries,
            template_path=str(template_path) if template_path else "",
            output_path=str(output_path),
        )
        document_result = self.document_export_service.export(export_input)

        # Add to export history
        history_item = ExportHistoryItem(
            output_path=document_result.output_path,
            exported_at=datetime.now(timezone.utc).isoformat(),
            reference_count=len(export_input.references),
            log_path=str(log_path.resolve()),
        )
        state.export.history.insert(0, history_item)
        # Keep only last 10 exports in history
        if len(state.export.history) > 10:
            state.export.history = state.export.history[:10]
        
        state.export.last_docx_path = document_result.output_path
        state.export.last_exported_at = datetime.now(timezone.utc).isoformat()
        state.export.status = GenericStatus.COMPLETED
        state.workflow_stage = WorkflowStage.DONE
        self.workspace_manager.save_state(project_id, state)

        log_payload = {
            "project_id": project_id,
            "exported_at": state.export.last_exported_at,
            "output_path": document_result.output_path,
            "template_path": export_input.template_path,
            "reference_count": len(export_input.references),
            "citations": export_input.citation_mapping,
        }
        log_path.write_text(json.dumps(log_payload, ensure_ascii=False, indent=2), encoding="utf-8")

        return ProjectExportResult(
            output_path=str(Path(document_result.output_path).resolve()),
            log_path=str(log_path.resolve()),
            reference_count=len(export_input.references),
            message=(
                "论文文档已导出。"
                if export_input.template_path
                else "论文文档已导出，当前使用内置基础样式；如提供学校模板，后续可直接切换为模板注入模式。"
            ),
        )

    def _validate_export_ready(self, state: ProjectState) -> None:
        if not state.chunk_plan.confirmed_plan:
            raise ConflictError("confirmed chunk plan is required before export")
        if not state.references.processed_items:
            raise ConflictError("at least one processed PDF is required before export")
        if not state.generation.blocks:
            raise ConflictError("generated blocks are required before export")
        if any(block.normalized_json is None for block in state.generation.blocks):
            raise ConflictError("all block contents must be imported before export")
        if state.generation.abstract_json is None:
            raise ConflictError("abstract must be imported before export")

    def _build_citation_mapping(self, state: ProjectState) -> dict[str, int]:
        mapping: dict[str, int] = {}
        for item in state.references.processed_items:
            mapping[str(item.effective_index)] = item.effective_index
            mapping[str(item.effective_index).zfill(2)] = item.effective_index
        return mapping

    def _resolve_template_path(self, template_id: str) -> Path | None:
        assets_root = Path(__file__).resolve().parents[3] / "assets" / "templates"
        template_path = assets_root / f"{template_id}.docx"
        return template_path if template_path.exists() else None

    def _build_abstract_block(self, state: ProjectState) -> dict[str, Any] | None:
        abstract_json = state.generation.abstract_json
        if not abstract_json:
            return None
        content: list[dict[str, Any]] = [{"type": "h1", "text": abstract_json.get("title", "摘要")}]
        for paragraph in abstract_json.get("content", []):
            text = str(paragraph).strip()
            if text:
                content.append({"type": "p", "text": text})
        keywords = [item.strip() for item in abstract_json.get("keywords", []) if str(item).strip()]
        if keywords:
            content.append({"type": "p", "text": f"关键词：{'；'.join(keywords)}"})
        return {
            "block_index": 0,
            "block_title": abstract_json.get("title", "摘要"),
            "content": content,
        }
